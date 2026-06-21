from pathlib import Path
from datetime import datetime
from docx import Document


class MarketingAgent:

    def _write_docx(self, path: Path, title: str, lines: list[str]):
        doc = Document()
        doc.add_heading(title, level=1)

        for line in lines:
            if line.strip() == "":
                doc.add_paragraph("")
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

        doc.save(path)

    def create_day11_content(self):
        base_dir = Path("outputs/marketing/day11")
        assets_dir = base_dir / "assets"

        base_dir.mkdir(parents=True, exist_ok=True)
        assets_dir.mkdir(parents=True, exist_ok=True)

        today = datetime.now().strftime("%Y-%m-%d")

        contents = {
            "facebook.docx": (
                "Day 11 - Facebook Post",
                [
                    f"Date: {today}",
                    "",
                    "Chủ đề: HGPT_AI_OS đã boot thành công Core v0.1",
                    "",
                    "Hôm nay HGPT chính thức hoàn thành bước đầu tiên của hành trình xây dựng hệ điều hành AI nội bộ: HGPT_AI_OS Core v0.1.",
                    "",
                    "Hệ thống đã chạy được:",
                    "- CLI",
                    "- Runtime",
                    "- Kernel",
                    "- Agent",
                    "- Task Create / List / Run / Complete",
                    "- Smoke Test PASS",
                    "",
                    "Đây là nền móng để HGPT từng bước tự động hóa quản lý, sản xuất, bảo trì, dự án, QA/QC và marketing.",
                    "",
                    "Mỗi ngày một bước nhỏ. Mỗi bước nhỏ tạo thành một nhà máy số.",
                    "",
                    "#HGPT #HGPTSteel #AIOS #DigitalFactory #Automation #SteelStructure",
                ],
            ),
            "tiktok.docx": (
                "Day 11 - TikTok Script",
                [
                    "Hook:",
                    "Chúng tôi vừa boot thành công hệ điều hành AI đầu tiên cho nhà máy cơ khí kết cấu thép.",
                    "",
                    "Body:",
                    "Đây là HGPT_AI_OS Core v0.1.",
                    "Nó có thể chạy lệnh, tạo task, lưu task, chạy task và hoàn tất task.",
                    "",
                    "Ứng dụng tiếp theo:",
                    "- Bảo trì",
                    "- QA/QC",
                    "- Quản lý dự án",
                    "- Sản xuất",
                    "- Marketing",
                    "",
                    "CTA:",
                    "Theo dõi hành trình xây dựng HGPT Steel Digital Factory mỗi ngày.",
                    "",
                    "#HGPT #AIOS #DigitalFactory #SteelFactory #Automation",
                ],
            ),
            "image_prompt.docx": (
                "Day 11 - Image Prompt",
                [
                    "A futuristic steel structure factory control room, digital dashboard, AI operating system interface, industrial steel fabrication environment, engineers monitoring automation workflows, clean modern lighting, professional corporate style, HGPT Steel Digital Factory concept.",
                ],
            ),
            "video_prompt.docx": (
                "Day 11 - Video Prompt",
                [
                    "Create a 15-second vertical video showing:",
                    "- A steel fabrication factory.",
                    "- A digital AI operating system dashboard booting up.",
                    "- Task flow: Create → Run → Complete.",
                    "- Engineers reviewing automation results.",
                    "- Closing text: HGPT_AI_OS Core v0.1 - Boot Passed.",
                    "",
                    "Style: modern industrial, professional, cinematic, high-tech factory automation.",
                ],
            ),
            "seo.docx": (
                "Day 11 - SEO",
                [
                    "Title: HGPT_AI_OS Core v0.1 chính thức boot thành công",
                    "",
                    "Meta Description:",
                    "HGPT Steel bắt đầu hành trình xây dựng hệ điều hành AI nội bộ phục vụ tự động hóa nhà máy cơ khí kết cấu thép.",
                    "",
                    "Keywords:",
                    "HGPT, HGPT Steel, AI OS, Digital Factory, Steel Structure Automation, Nhà máy số, Tự động hóa cơ khí",
                ],
            ),
            "approval_checklist.docx": (
                "Day 11 - Approval Checklist",
                [
                    "- Đọc Facebook post",
                    "- Đọc TikTok script",
                    "- Kiểm tra hashtag",
                    "- Kiểm tra image prompt",
                    "- Kiểm tra video prompt",
                    "- Duyệt đăng Facebook",
                    "- Duyệt đăng TikTok",
                ],
            ),
        }

        for filename, (title, lines) in contents.items():
            self._write_docx(base_dir / filename, title, lines)

        return base_dir
