from pathlib import Path
import re

from docx import Document

from hgpt_ai_os.diagnostics import instrument_runtime_tracing, module_loaded, trace_call


class DocxExporter:
    _HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
    _BULLET_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
    _BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
    _FORBIDDEN_TEMPLATE_LABEL_RE = re.compile(
        r"(?im)^\s*"
        r"("
        r"Story|Pain|Truth|Hook|Curiosity|Practical actions|"
        r"Scene|Camera|Voice|Caption|Transition|Negative prompt|"
        r"Duration|Aspect ratio|Lens|Lighting|Composition|Mood|Texture|"
        r"Problem|One practical tip|CTA"
        r")\s*:",
    )
    _FORBIDDEN_LEAKAGE_PHRASE_RE = re.compile(
        r"\b("
        r"Manager's job|Quality release|Hold the product|The topic belongs to|"
        r"Root cause analysis|Root Causes|"
        r"Call To Action|Lesson learned|Real shop scenario|Practical solution"
        r")\b",
        re.IGNORECASE,
    )
    _BROKEN_VIETNAMESE_RE = re.compile(
        r"\b("
        r"[Đđ]\s+ng\s+c|"
        r"ki\s+m\s+tra|"
        r"b\s+ng\s+ch\s+ng"
        r")\b",
        re.IGNORECASE,
    )

    def save(self, path: Path, title: str, content: str):
        trace_call("DOCX Writer.save", self, selected_topic=title, output_file=path)

        path.parent.mkdir(parents=True, exist_ok=True)
        self.validate_content(str(content))

        doc = Document()
        self._set_vietnamese_font(doc)
        doc.add_heading(title, level=1)
        self._add_markdown_content(doc, str(content))
        doc.save(path)
        self.validate_rendered_docx(path)
        trace_call("DOCX save completed", self, selected_topic=title, output_file=path, final_docx_writer=self.__class__.__name__)

    def validate_content(self, content: str) -> None:
        if not content.strip():
            raise ValueError("DOCX export blocked: empty content.")
        if content.encode("utf-8").decode("utf-8") != content:
            raise ValueError("DOCX export blocked: invalid UTF-8 content.")
        if self._has_duplicate_paragraphs(content):
            raise ValueError("DOCX export blocked: duplicated paragraphs.")
        forbidden = self._FORBIDDEN_TEMPLATE_LABEL_RE.search(content)
        if forbidden:
            label = forbidden.group(1)
            raise ValueError(
                f"DOCX export blocked: forbidden English text '{label}'."
            )
        forbidden = self._FORBIDDEN_LEAKAGE_PHRASE_RE.search(content)
        if forbidden:
            raise ValueError(
                f"DOCX export blocked: forbidden English text '{forbidden.group(0)}'."
            )
        broken = self._BROKEN_VIETNAMESE_RE.search(content)
        if broken:
            raise ValueError(
                f"DOCX export blocked: broken Vietnamese text '{broken.group(0)}'."
            )

    def validate_rendered_docx(self, path: Path) -> None:
        try:
            doc = Document(path)
        except Exception as exc:
            raise ValueError(f"DOCX export blocked: incomplete rendering: {exc}") from exc
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()
        if not text:
            raise ValueError("DOCX export blocked: rendered DOCX is empty.")

    def _has_duplicate_paragraphs(self, content: str) -> bool:
        paragraphs = [
            re.sub(r"\s+", " ", paragraph.strip().lower())
            for paragraph in re.split(r"\n{2,}", content)
            if len(paragraph.strip()) > 80
        ]
        return len(paragraphs) != len(set(paragraphs))

    def _set_vietnamese_font(self, doc: Document) -> None:
        for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
            style = doc.styles[style_name]
            style.font.name = "Arial"

    def _add_markdown_content(self, doc: Document, content: str):
        for raw_line in content.splitlines():
            line = raw_line.strip()

            if not line:
                doc.add_paragraph()
                continue

            heading = self._HEADING_RE.match(line)
            if heading:
                level = min(len(heading.group(1)), 3)
                paragraph = doc.add_heading(level=level)
                self._add_inline_markdown(paragraph, heading.group(2).strip())
                continue

            bullet = self._BULLET_RE.match(line)
            if bullet:
                paragraph = doc.add_paragraph(style="List Bullet")
                self._add_inline_markdown(paragraph, bullet.group(1).strip())
                continue

            paragraph = doc.add_paragraph()
            self._add_inline_markdown(paragraph, line)

    def _add_inline_markdown(self, paragraph, text: str):
        position = 0

        for match in self._BOLD_RE.finditer(text):
            if match.start() > position:
                paragraph.add_run(text[position : match.start()])

            run = paragraph.add_run(match.group(1))
            run.bold = True
            position = match.end()

        if position < len(text):
            paragraph.add_run(text[position:])


instrument_runtime_tracing(globals())
module_loaded(__name__, __file__, DocxExporter)
