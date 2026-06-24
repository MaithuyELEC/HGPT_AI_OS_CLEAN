
93
from pathlib import Path
from docx import Document

from hgpt_ai_os.planner.planner_engine import PlannerEngine

from hgpt_ai_os.knowledge.retriever import KnowledgeRetriever
from hgpt_ai_os.knowledge.bundle import KnowledgeBundle

from hgpt_ai_os.content.generator import ContentGenerator


class LucidOrchestrator:

    def run(self):

        planner = PlannerEngine()

        task = planner.next_task()

        if not task:
            print("No TODO task.")
            return

        planner.update_status(task["row"], "RUNNING")

        try:

            retriever = KnowledgeRetriever()

        items = retriever.retrieve(
            task["topic"],
            top_k=5
        )

        bundle = KnowledgeBundle(
            query=task["topic"],
            items=items
        )

        context = bundle.context()

        generator = ContentGenerator()

        facebook = generator.generate_facebook(
            task["topic"],
            context
        )

        seo = generator.generate_seo(
            task["topic"],
            context
        )

        out = Path("outputs/marketing") / f"Day{int(task['day']):03d}"

        out.mkdir(
            parents=True,
            exist_ok=True
        )

        files = {
            "facebook.docx": facebook,
            "seo.docx": seo,
        }

        for name, content in files.items():

            doc = Document()

            doc.add_heading(
                task["topic"],
                level=1
            )

            doc.add_paragraph(content)

            doc.save(out / name)

        planner.update_status(
            task["row"],
            "GENERATED"
        )

        except Exception:

            planner.update_status(
                task["row"],
                "TODO"
            )

            raise

        print("=" * 60)
        print("HGPT LUCID AUTO")
        print("=" * 60)
        print("Task :", task["topic"])
        print("Knowledge :", len(items))
        print("Status : GENERATED")
        print("Output :", out)


if __name__ == "__main__":
    LucidOrchestrator().run()
