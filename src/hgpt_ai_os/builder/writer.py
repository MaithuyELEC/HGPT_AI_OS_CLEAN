from pathlib import Path
from docx import Document


class DocxWriter:
    def write(self, output_file, title, paragraphs):

        output_file = Path(output_file)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        doc.add_heading(str(title), level=1)

        for p in paragraphs:
            doc.add_paragraph(str(p))

        doc.save(output_file)

        return output_file