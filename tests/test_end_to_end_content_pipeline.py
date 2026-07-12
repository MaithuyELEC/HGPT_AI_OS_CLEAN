from __future__ import annotations

import io
import os
import re
import sys
import tempfile
import unittest
from contextlib import redirect_stdout
from pathlib import Path
from unittest.mock import patch

from docx import Document


SRC = Path(__file__).resolve().parents[1] / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

REQUIRED_DOCX_FILES = (
    "approval_checklist.docx",
    "facebook.docx",
    "hashtags.docx",
    "image_prompt.docx",
    "seo.docx",
    "tiktok.docx",
    "video_prompt.docx",
)

STEEL_TOPICS = (
    ("AWS D1.1", ("aws d1.1", "hàn", "welding", "steel", "thép")),
    ("Lỗi Fit-up", ("fit-up", "hàn", "welding", "steel", "thép")),
)


class EndToEndContentPipelineTest(unittest.TestCase):
    def test_release_topics_generate_docx_with_topic_appropriate_knowledge(self):
        from hgpt_ai_os import production

        previous_cwd = Path.cwd()
        repo_root = Path(__file__).resolve().parents[1]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output_root = root / "outputs"
            profile = root / "profile"
            free_env = {
                "USERPROFILE": str(profile),
                "AI_PROVIDER": "none",
                "OPENAI_API_KEY": "",
                "GEMINI_API_KEY": "",
                "GOOGLE_API_KEY": "",
                "ANTHROPIC_API_KEY": "",
            }

            try:
                os.chdir(repo_root)
                with (
                    patch.dict(os.environ, free_env, clear=True),
                    patch.object(production, "OUTPUT_ROOT", output_root),
                    patch(
                        "urllib.request.urlopen",
                        side_effect=AssertionError("remote provider call attempted"),
                    ),
                    redirect_stdout(io.StringIO()),
                ):
                    generated_text = {
                        topic: self._generate_and_read_docx_text(production, day, topic)
                        for day, (topic, _terms) in enumerate(
                            STEEL_TOPICS,
                            start=1,
                        )
                    }
            finally:
                os.chdir(previous_cwd)

        for topic, expected_terms in STEEL_TOPICS:
            body = generated_text[topic]
            self.assertTrue(
                any(self._contains_term(body, term) for term in expected_terms),
                f"{topic!r} should include steel knowledge in generated DOCX text.",
            )

    def _generate_and_read_docx_text(self, production, day: int, topic: str) -> str:
        output_dir = production.build_outputs(day, topic, open_output_folder=False)
        self.assertTrue(output_dir.exists())
        self.assertEqual(
            sorted(path.name for path in output_dir.glob("*.docx")),
            list(REQUIRED_DOCX_FILES),
        )

        return "\n".join(
            self._read_docx_text(path)
            for path in sorted(output_dir.glob("*.docx"))
        ).lower()

    def _read_docx_text(self, path: Path) -> str:
        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(paragraph.text for paragraph in cell.paragraphs)
        return "\n".join(parts)

    def _contains_term(self, body: str, term: str) -> bool:
        pattern = r"(?<!\w)" + re.escape(term.lower()) + r"(?!\w)"
        return re.search(pattern, body) is not None


if __name__ == "__main__":
    unittest.main()
