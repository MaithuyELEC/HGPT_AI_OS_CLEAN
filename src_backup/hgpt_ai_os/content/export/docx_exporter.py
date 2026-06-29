from pathlib import Path
from docx import Document


class DocxExporter:

    def save(self, path: Path, title: str, content: str):

        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(str(content))
        doc.save(path)
